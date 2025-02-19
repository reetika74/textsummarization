import streamlit as st
from googletrans import Translator, LANGUAGES
from io import BytesIO
from docx import Document
from transformers import pipeline

# Initialize the translator
translator = Translator()

# Initialize the summarization pipeline
summarizer = pipeline("summarization", model="t5-small", tokenizer="t5-small", framework="pt")

# Function to translate text
def translate_text(text, dest_language='en'):
    try:
        translation = translator.translate(text, dest=dest_language)
        return translation.text
    except Exception as e:
        return f"Translation error: {e}"

# Function to summarize text
def summarize_text(text, max_length=130, min_length=30):
    try:
        summary = summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
        return summary[0]['summary_text']
    except Exception as e:
        return f"Summarization error: {e}"

# App title
st.title("The Summarize Pro")

# Sidebar for summarization settings
with st.sidebar:
    st.header("Summarization Settings")
    
    # Summarization method selection
    summary_type = st.radio("Choose summarization method:", ('Extractive', 'Abstractive'))

    # Slider to select summary length
    summary_length = st.slider("Summary length (in percentage of original text):", 10, 100, 50)

    # Language selection for input text
    input_language_full_names = list(LANGUAGES.values())
    input_language_codes = list(LANGUAGES.keys())
    input_language_selection = st.selectbox("Select input language:", ['Select an input language'] + input_language_full_names)

    # Language selection for output summary
    output_language_full_names = list(LANGUAGES.values())
    output_language_codes = list(LANGUAGES.keys())
    output_language_selection = st.selectbox("Select output language:", ['Select an output language'] + output_language_full_names)

    # Find the corresponding language codes if a valid option is selected
    if input_language_selection != 'Select an input language':
        input_language = input_language_codes[input_language_full_names.index(input_language_selection)]
    else:
        input_language = None  # No valid input language selected

    if output_language_selection != 'Select an output language':
        output_language = output_language_codes[output_language_full_names.index(output_language_selection)]
    else:
        output_language = None  # No valid output language selected

# Input text box
text_input = st.text_area("Paste your text here", height=200)

# Option to upload .txt and .docx files only
uploaded_file = st.file_uploader("Or upload a file", type=['txt', 'docx'])

# Process button
if st.button("Summarize"):
    if (text_input or uploaded_file) and input_language and output_language:
        # Read uploaded file content
        if uploaded_file:
            if uploaded_file.type == "text/plain":
                text_input = uploaded_file.getvalue().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                text_input = "\n".join([para.text for para in doc.paragraphs])

        # Translate input text to English if it’s in a different language
        if input_language != 'en':
            translated_input = translate_text(text_input, dest_language='en')
        else:
            translated_input = text_input

        # Summarization logic
        # Adjust summary length parameters based on user input
        max_length = int(len(translated_input.split()) * (summary_length / 100))
        min_length = max(10, max_length // 2)  # Ensure a minimum summary length
        summary = summarize_text(translated_input, max_length=max_length, min_length=min_length)

        # Translate summary to the selected output language
        if output_language != 'en':
            summary = translate_text(summary, dest_language=output_language)

        st.write("### Summary:")
        st.write(summary)

        # Add download options for .txt and .docx
        if summary:
            # Download as .txt
            st.download_button(
                label="Download as .txt",
                data=summary,
                file_name='summary.txt',
                mime='text/plain'
            )

            # Download as .docx
            doc = Document()
            doc.add_paragraph(summary)
            byte_io = BytesIO()
            doc.save(byte_io)
            byte_io.seek(0)

            st.download_button(
                label="Download as .docx",
                data=byte_io,
                file_name='summary.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    else:
        st.warning("Please provide text or upload a file, and select both input and output languages to summarize.")
